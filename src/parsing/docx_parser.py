"""
Clause-aware parser for 3GPP .docx specifications.

3GPP specs are written against a standard Word template, so paragraph
*styles* (not visual formatting) reliably tell us the document structure:

    Heading 1..Heading 9   clause / sub-clause headings, e.g. "4.2.8.1\tGeneral"
    toc 1..toc 9, TT       table of contents entries -- not real content, skipped
    Normal                 body text
    NO                     NOTE
    EX                     EXAMPLE
    B1..B4                 bulleted list items (nesting level 1-4)
    TH / TF                table caption ("Table 5.1-1: ...") / table footnote

We walk the document body in original order (paragraphs AND tables
interleaved, since python-docx exposes them as separate collections by
default) so that a table stays attached to the clause it actually appears
in, and split the document into "sections": one section per heading, holding
every block of content until the next heading of any level. Each section
carries its full clause number and heading breadcrumb, which is what lets
the RAG layer answer "per TS 23.501 clause 4.2.8.1" instead of a vague
paraphrase -- the citation-groundedness that keeps hallucinations down.

Deliberately NOT handled here: cross-reference resolution ("see clause
5.4.2"), figures/images, revision-mark diffs. Out of scope for ingestion v1.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

import docx
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

HEADING_STYLE_RE = re.compile(r"^Heading (\d)$")
TOC_STYLE_RE = re.compile(r"^(toc \d|TT)$", re.IGNORECASE)
NOTE_STYLE_RE = re.compile(r"^(NO|NO1|NO2|NO3)$")
EXAMPLE_STYLE_RE = re.compile(r"^EX$")
BULLET_STYLE_RE = re.compile(r"^B(\d)$")
TABLE_CAPTION_STYLE_RE = re.compile(r"^TH$")

# "4.2.8.1\tGeneral", "D.1\tIntroduction", "K.2.2.3\tConfiguration...",
# "4.2.5a\tRadio Capabilities...", "4.2.8.1A\tGeneral Concepts..."
# First component is either a bare annex letter ("D") or digits with an
# optional trailing letter suffix ("4", "5a"); later components are always
# dot-prefixed digits with an optional trailing letter suffix.
NUMBERED_HEADING_RE = re.compile(
    r"^((?:[A-Z]|\d+[A-Za-z]?)(?:\.\d+[A-Za-z]?)*)\t(.+)$", re.DOTALL
)
# "Annex D (informative):\n5GS support for ..."
ANNEX_HEADING_RE = re.compile(r"^Annex\s+([A-Z]+)\s*\(([^)]+)\):\s*\n?(.*)$", re.DOTALL)


@dataclass
class Block:
    type: str  # "paragraph" | "note" | "example" | "bullet" | "table"
    text: str = ""
    bullet_level: int | None = None
    table_caption: str | None = None
    table_rows: list[list[str]] = field(default_factory=list)


@dataclass
class Section:
    clause_number: str | None   # None for unnumbered headings like "Foreword"
    title: str
    level: int
    heading_path: str            # "4 > 4.2 > 4.2.8 > 4.2.8.1 General Concepts..."
    blocks: list[Block] = field(default_factory=list)

    def text(self) -> str:
        parts = []
        for b in self.blocks:
            if b.type == "table":
                cap = f"{b.table_caption}\n" if b.table_caption else ""
                rows_txt = "\n".join(" | ".join(row) for row in b.table_rows)
                parts.append(f"{cap}{rows_txt}")
            else:
                parts.append(b.text)
        return "\n".join(p for p in parts if p.strip())


@dataclass
class ParsedSpec:
    spec_id: str
    title: str
    release: str
    version: str
    sections: list[Section]


def _iter_block_items(document: DocumentObject) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables in the order they appear in the document body."""
    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _table_to_rows(table: Table) -> list[list[str]]:
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        if any(cells):
            rows.append(cells)
    return rows


def _parse_heading(text: str) -> tuple[str | None, str]:
    """Return (clause_number, title) for a heading paragraph's text."""
    text = text.strip()
    m = ANNEX_HEADING_RE.match(text)
    if m:
        letter, category, title = m.groups()
        return f"Annex {letter}", f"({category}) {title.strip()}"
    m = NUMBERED_HEADING_RE.match(text)
    if m:
        number, title = m.groups()
        return number, title.strip()
    return None, text.replace("\n", " ").strip()


def _clause_level(clause_number: str | None, style_level: int) -> int:
    if clause_number is None:
        return style_level
    if clause_number.startswith("Annex"):
        return 1
    return clause_number.count(".") + 1


def parse_docx(path: Path, spec_id: str) -> ParsedSpec:
    document = docx.Document(str(path))
    title, release, version = _extract_metadata(document, spec_id)

    sections: list[Section] = []
    heading_stack: list[str] = []  # breadcrumb of "clause title" strings by level
    current: Section | None = None
    pending_table_caption: str | None = None

    for item in _iter_block_items(document):
        if isinstance(item, Table):
            rows = _table_to_rows(item)
            if not rows:
                continue
            block = Block(type="table", table_caption=pending_table_caption, table_rows=rows)
            pending_table_caption = None
            if current is not None:
                current.blocks.append(block)
            continue

        para: Paragraph = item
        style_name = para.style.name if para.style else "Normal"
        text = para.text

        if TOC_STYLE_RE.match(style_name):
            continue
        if not text.strip():
            continue

        heading_m = HEADING_STYLE_RE.match(style_name)
        if heading_m or style_name.upper() in {"H6", "H7", "H8", "H9"}:
            style_level = int(heading_m.group(1)) if heading_m else int(style_name[1])
            clause_number, clause_title = _parse_heading(text)
            level = _clause_level(clause_number, style_level)

            heading_stack = heading_stack[: level - 1]
            crumb = f"{clause_number} {clause_title}" if clause_number else clause_title
            heading_stack.append(crumb.strip())

            current = Section(
                clause_number=clause_number,
                title=clause_title,
                level=level,
                heading_path=" > ".join(heading_stack),
            )
            sections.append(current)
            continue

        if TABLE_CAPTION_STYLE_RE.match(style_name):
            pending_table_caption = text.strip()
            continue

        if current is None:
            # Content before the first real heading (e.g. cover page fragments).
            continue

        if NOTE_STYLE_RE.match(style_name):
            current.blocks.append(Block(type="note", text=text.strip()))
        elif EXAMPLE_STYLE_RE.match(style_name):
            current.blocks.append(Block(type="example", text=text.strip()))
        elif BULLET_STYLE_RE.match(style_name):
            lvl = int(BULLET_STYLE_RE.match(style_name).group(1))
            current.blocks.append(Block(type="bullet", text=text.strip(), bullet_level=lvl))
        else:
            current.blocks.append(Block(type="paragraph", text=text.strip()))

    return ParsedSpec(spec_id=spec_id, title=title, release=release, version=version, sections=sections)


def _extract_metadata(document: DocumentObject, spec_id: str) -> tuple[str, str, str]:
    # The cover page layout isn't consistent across 3GPP's own templates: some
    # specs put "3GPP TS x V17.y.z" in a table cell, others in a plain
    # paragraph (styles "ZA"/"ZB"/"ZT"). Search both rather than assume one.
    version_re = re.compile(rf"3GPP T[SR] {re.escape(spec_id)} V(\d+)\.(\d+)\.(\d+)")

    def find_in(texts: Iterator[str]) -> tuple[str, str] | None:
        for text in texts:
            m = version_re.search(text)
            if m:
                return m.group(1), f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
        return None

    cell_texts = (cell.text for table in document.tables for row in table.rows for cell in row.cells)
    found = find_in(cell_texts)
    if found is None:
        found = find_in(p.text for p in document.paragraphs[:20])

    release, version = found if found else ("unknown", "unknown")

    title = document.core_properties.subject or document.core_properties.title or spec_id
    title = re.sub(r"\s*\(Release\s*\d+\)\s*$", "", title).strip()
    return title, release, version
